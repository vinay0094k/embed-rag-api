import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Set

from app.rag_components.models import RawDocument

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: Set[str] = {
    ".pdf", ".docx", ".md", ".markdown",
    ".txt", ".log",
    ".py", ".js", ".ts", ".go", ".rs", ".java",
    ".sh", ".bash",
    ".yaml", ".yml", ".tf", ".tfvars", ".hcl",
    ".json", ".dockerfile",
}

TEXT_EXTENSIONS: Set[str] = SUPPORTED_EXTENSIONS - {".pdf", ".docx"}

BINARY_EXTENSIONS: Set[str] = {".pdf", ".docx"}

BLOCKED_FILENAMES: Set[str] = {
    ".env", ".env.local", ".env.*.local",
    ".secrets", ".credentials", "credentials.json",
    ".ssh", "id_rsa", "id_ed25519",
    "docker-compose.override.yml",
}


class DocumentLoader:
    def __init__(self, data_root: Optional[str] = None):
        self.data_root = Path(data_root) if data_root else None

    @property
    def supported_extensions(self) -> Set[str]:
        return SUPPORTED_EXTENSIONS

    def load_file(self, file_path: str) -> Optional[RawDocument]:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            logger.warning("File not found: %s", file_path)
            return None

        if path.name in BLOCKED_FILENAMES or path.name.startswith(".env"):
            logger.warning("Blocked file (security): %s", file_path)
            return None

        ext = path.suffix.lower()
        _name_lower = path.name.lower()
        if ext == ".dockerfile" or "dockerfile" in _name_lower or _name_lower == "dockerfile":
            ext = ".dockerfile"

        if ext not in SUPPORTED_EXTENSIONS:
            logger.debug("Unsupported extension: %s (%s)", file_path, ext)
            return None

        stat = path.stat()
        logger.debug("Loading file: %s (ext=%s, size=%d)", file_path, ext, stat.st_size)
        doc_modified = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
        doc_created = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat()

        if ext == ".pdf":
            content = self._load_pdf(path)
        elif ext == ".docx":
            content = self._load_docx(path)
        else:
            content = self._load_text(path)

        if content is None:
            logger.warning("Failed to extract content from: %s", file_path)
            return None

        logger.info("Loaded file: %s (ext=%s, size=%d)", file_path, ext, stat.st_size)
        return RawDocument(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_extension=ext,
            content=content,
            file_size=stat.st_size,
            doc_modified_at=doc_modified,
            doc_created_at=doc_created,
        )

    def load_directory(
        self,
        dir_path: str,
        recursive: bool = True,
        glob_pattern: Optional[str] = None,
    ) -> List[RawDocument]:
        path = Path(dir_path)
        if not path.exists() or not path.is_dir():
            logger.warning("Directory not found: %s", dir_path)
            return []

        documents: List[RawDocument] = []

        if glob_pattern:
            files = list(path.rglob(glob_pattern)) if recursive else list(path.glob(glob_pattern))
        else:
            files = []
            if recursive:
                for ext in SUPPORTED_EXTENSIONS:
                    files.extend(path.rglob(f"*{ext}"))
                files.extend(path.rglob("Dockerfile"))
            else:
                for ext in SUPPORTED_EXTENSIONS:
                    files.extend(path.glob(f"*{ext}"))
                files.extend(path.glob("Dockerfile"))

            files = list(set(files))

        for file_path in sorted(files):
            if not file_path.is_file():
                continue
            if file_path.name.startswith("."):
                continue
            if file_path.name in BLOCKED_FILENAMES or file_path.name.startswith(".env"):
                logger.debug("Skipping blocked file: %s", file_path)
                continue

            doc = self.load_file(str(file_path))
            if doc is not None:
                documents.append(doc)

        logger.info("Loaded directory: %s — %d documents", dir_path, len(documents))
        return documents

    def load_files(self, file_paths: List[str]) -> List[RawDocument]:
        documents: List[RawDocument] = []
        for fp in file_paths:
            doc = self.load_file(fp)
            if doc is not None:
                documents.append(doc)
        logger.info("Loaded %d of %d requested files", len(documents), len(file_paths))
        return documents

    def _load_text(self, path: Path) -> Optional[str]:
        try:
            logger.debug("Reading text file (utf-8): %s", path)
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            logger.debug("UTF-8 read failed for %s: %s", path, e)
            try:
                logger.debug("Retrying with latin-1: %s", path)
                return path.read_text(encoding="latin-1", errors="replace")
            except Exception as e2:
                logger.error("Failed to read text file: %s — %s", path, e2)
                return None

    def _load_pdf(self, path: Path) -> Optional[str]:
        try:
            import fitz
        except ImportError:
            raise ImportError("pymupdf is required for PDF loading: pip install pymupdf")

        try:
            logger.debug("Loading PDF: %s", path)
            doc = fitz.open(str(path))
            pages: List[str] = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            doc.close()

            if not pages:
                logger.debug("No text extracted from PDF, trying OCR: %s", path)
                try:
                    doc = fitz.open(str(path))
                    for page in doc:
                        tp = page.get_textpage_ocr(
                            full=True, language="eng"
                        )
                        text = tp.extractText() or ""
                        if text.strip():
                            pages.append(text)
                        tp.close()
                    doc.close()
                    if pages:
                        logger.info("Loaded PDF via OCR: %s (%d pages)", path, len(pages))
                        return "\n\n--- Page Break ---\n\n".join(pages)
                except Exception as e:
                    logger.warning("OCR failed for %s: %s", path, e)
                logger.warning("No content extracted from PDF: %s", path)
                return None

            logger.info("Loaded PDF: %s (%d pages)", path, len(pages))
            return "\n\n--- Page Break ---\n\n".join(pages)
        except Exception as e:
            logger.error("Failed to load PDF: %s — %s", path, e)
            return None

    def _load_docx(self, path: Path) -> Optional[str]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading: pip install python-docx")

        try:
            logger.debug("Loading DOCX: %s", path)
            doc = Document(str(path))
            paragraphs: List[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            tables: List[str] = []
            for table in doc.tables:
                rows: List[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables.append("\n".join(rows))

            parts: List[str] = []
            if paragraphs:
                parts.append("\n\n".join(paragraphs))
            if tables:
                parts.append("\n\n--- Table ---\n\n" + "\n\n".join(tables))

            content = "\n\n".join(parts) if parts else None
            if content:
                logger.info("Loaded DOCX: %s (%d chars)", path, len(content))
            else:
                logger.warning("No content extracted from DOCX: %s", path)
            return content
        except Exception as e:
            logger.error("Failed to load DOCX: %s — %s", path, e)
            return None
